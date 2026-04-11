"""
Generate enterprise reference.docx template for Pandoc.

Defines corporate-styled Word styles that Pandoc maps to automatically:
Heading 1-4, Normal, Source Code, Block Text, Table Grid.

Usage:
    pip install python-docx
    python scripts/generate-reference-docx.py

Output:
    apps/desktop/src-tauri/resources/pandoc/reference.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Enterprise color palette (Office-compatible)
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE = RGBColor(0x2E, 0x74, 0xB5)
ACCENT_BLUE = RGBColor(0x44, 0x72, 0xC4)
DARK_GRAY = RGBColor(0x44, 0x54, 0x6A)
TEXT_COLOR = RGBColor(0x1F, 0x29, 0x37)
LIGHT_GRAY_BG = RGBColor(0xF5, 0xF5, 0xF5)


def setup_page(doc):
    """Configure page layout: Letter, 1 inch margins."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def setup_header_footer(doc):
    """Add page numbers in footer."""
    section = doc.sections[0]

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page number field
    run = footer_para.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = footer_para.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)

    run3 = footer_para.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)

    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = DARK_GRAY


def style_heading(style, font_name, size, color, bold=False, space_before=0, space_after=0):
    """Configure a heading style."""
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.color.rgb = color
    font.bold = bold

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.keep_with_next = True


def setup_styles(doc):
    """Define enterprise styles matching Pandoc's expected style names."""
    styles = doc.styles

    # --- Heading 1 ---
    h1 = styles['Heading 1']
    style_heading(h1, 'Calibri Light', 26, DARK_BLUE, space_before=24, space_after=6)

    # --- Heading 2 ---
    h2 = styles['Heading 2']
    style_heading(h2, 'Calibri Light', 20, MED_BLUE, space_before=18, space_after=4)

    # --- Heading 3 ---
    h3 = styles['Heading 3']
    style_heading(h3, 'Calibri', 16, MED_BLUE, space_before=12, space_after=4)

    # --- Heading 4 ---
    h4 = styles['Heading 4']
    style_heading(h4, 'Calibri', 14, DARK_GRAY, bold=True, space_before=10, space_after=2)

    # --- Normal (body text) ---
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # --- First Paragraph (Pandoc uses this for first para after heading) ---
    if 'First Paragraph' not in [s.name for s in styles]:
        fp = styles.add_style('First Paragraph', WD_STYLE_TYPE.PARAGRAPH)
        fp.base_style = normal
        fp.font.name = 'Calibri'
        fp.font.size = Pt(11)
        fp.font.color.rgb = TEXT_COLOR

    # --- Source Code (for code blocks) ---
    if 'Source Code' not in [s.name for s in styles]:
        sc = styles.add_style('Source Code', WD_STYLE_TYPE.PARAGRAPH)
    else:
        sc = styles['Source Code']
    sc.font.name = 'Consolas'
    sc.font.size = Pt(9.5)
    sc.font.color.rgb = TEXT_COLOR
    sc.paragraph_format.space_before = Pt(4)
    sc.paragraph_format.space_after = Pt(4)
    sc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add light gray background shading to Source Code
    ppr = sc.element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
    ppr.append(shading)

    # --- Block Text (for blockquotes) ---
    if 'Block Text' not in [s.name for s in styles]:
        bt = styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    else:
        bt = styles['Block Text']
    bt.font.name = 'Calibri'
    bt.font.size = Pt(11)
    bt.font.italic = True
    bt.font.color.rgb = DARK_GRAY
    bt.paragraph_format.left_indent = Inches(0.5)
    bt.paragraph_format.space_before = Pt(6)
    bt.paragraph_format.space_after = Pt(6)

    # Left border for blockquote
    ppr = bt.element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="4472C4"/>'
        f'</w:pBdr>'
    )
    ppr.append(borders)

    # --- Verbatim Char (inline code) ---
    if 'Verbatim Char' not in [s.name for s in styles]:
        vc = styles.add_style('Verbatim Char', WD_STYLE_TYPE.CHARACTER)
    else:
        vc = styles['Verbatim Char']
    vc.font.name = 'Consolas'
    vc.font.size = Pt(10)
    vc.font.color.rgb = TEXT_COLOR

    # --- Table styles ---
    # Pandoc uses the default Table style; we style it via the document's default table
    # The actual table styling happens through Pandoc's built-in table handling


def add_sample_content(doc):
    """Add minimal sample content so Pandoc can extract all styles."""
    doc.add_heading('Sample Heading 1', level=1)
    doc.add_heading('Sample Heading 2', level=2)
    doc.add_heading('Sample Heading 3', level=3)
    doc.add_heading('Sample Heading 4', level=4)
    doc.add_paragraph('Sample body text in Normal style.')

    # Add Source Code paragraph
    style_names = [s.name for s in doc.styles]
    if 'Source Code' in style_names:
        doc.add_paragraph('def hello(): pass', style='Source Code')

    # Add Block Text paragraph
    if 'Block Text' in style_names:
        doc.add_paragraph('This is a blockquote.', style='Block Text')

    # Add First Paragraph
    if 'First Paragraph' in style_names:
        doc.add_paragraph('First paragraph after heading.', style='First Paragraph')


def main():
    doc = Document()

    setup_page(doc)
    setup_header_footer(doc)
    setup_styles(doc)
    add_sample_content(doc)

    # Output path
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    output_path = os.path.join(
        project_root,
        'apps', 'desktop', 'src-tauri', 'resources', 'pandoc', 'reference.docx'
    )

    doc.save(output_path)
    print(f"Generated: {output_path}")
    print(f"Size: {os.path.getsize(output_path)} bytes")


if __name__ == '__main__':
    main()
